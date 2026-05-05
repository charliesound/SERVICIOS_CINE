from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
import csv
import json
import mimetypes
import re
import zipfile
from xml.etree import ElementTree as ET

from fastapi import HTTPException
from pypdf import PdfReader
from docx import Document as DocxDocument
from sqlalchemy import and_, desc, select
from sqlalchemy.ext.asyncio import AsyncSession

from models.ingest_document import (
    DocumentAsset,
    DocumentClassification,
    DocumentExtraction,
    DocumentStructuredData,
)
from models.ingest_handshake import IngestEvent, StorageSource
from models.ingest_scan import MediaAsset
from services.ingest_scan_service import get_owned_asset
from services.script_document_classifier import is_probable_screenplay
from services.storage_handshake_service import (
    get_owned_storage_source,
    log_ingest_event,
    normalize_mounted_path,
    resolve_organization_id,
)


IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".heic"}
TEXT_EXTENSIONS = {".txt"}
PDF_EXTENSIONS = {".pdf"}
DOC_EXTENSIONS = {".doc"}
DOCX_EXTENSIONS = {".docx"}
CSV_EXTENSIONS = {".csv"}
XLS_EXTENSIONS = {".xls"}
XLSX_EXTENSIONS = {".xlsx"}

DOC_TYPE_RULES = {
    "sound_report": ["roll", "mixer", "boom", "sample rate"],
    "camera_report": ["camera", "mag", "card", "lens", "fps"],
    "script_note": ["continuity", "best take", "editor note"],
    "director_note": ["intention", "pacing", "coverage"],
    "operator_note": ["operator", "operator note", "camera operator", "sound operator"],
}


def detect_file_metadata(
    file_name: str, original_path: Optional[str]
) -> Tuple[str, Optional[str]]:
    extension = Path(file_name).suffix.lower()
    guessed_mime, _encoding = mimetypes.guess_type(original_path or file_name)
    return extension, guessed_mime


def get_document_source_path(
    document: DocumentAsset, media_asset: Optional[MediaAsset]
) -> str:
    if document.original_path:
        return normalize_mounted_path(document.original_path)
    if media_asset is not None:
        return normalize_mounted_path(media_asset.canonical_path)
    raise HTTPException(
        status_code=400, detail="Document has no resolvable source path"
    )


def read_text_file(path: str) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return Path(path).read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return Path(path).read_text(errors="ignore")


def extract_docx_text(path: str) -> str:
    document = DocxDocument(path)
    parts = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_pdf_text(path: str) -> str:
    reader = PdfReader(path)
    parts: List[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())
    return "\n\n".join(parts)


def extract_csv_table(path: str) -> Dict[str, Any]:
    with open(path, "r", encoding="utf-8-sig", newline="") as csv_file:
        reader = csv.DictReader(csv_file)
        rows = list(reader)
        return {
            "table_type": "csv",
            "headers": reader.fieldnames or [],
            "rows": rows,
        }


def parse_xlsx_shared_strings(archive: zipfile.ZipFile) -> List[str]:
    if "xl/sharedStrings.xml" not in archive.namelist():
        return []
    root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
    namespace = {"main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    values: List[str] = []
    for item in root.findall("main:si", namespace):
        text_parts = [node.text or "" for node in item.findall(".//main:t", namespace)]
        values.append("".join(text_parts))
    return values


def parse_xlsx_table(path: str) -> Dict[str, Any]:
    with zipfile.ZipFile(path) as archive:
        shared_strings = parse_xlsx_shared_strings(archive)
        sheet_names = sorted(
            [
                name
                for name in archive.namelist()
                if name.startswith("xl/worksheets/sheet") and name.endswith(".xml")
            ]
        )
        if not sheet_names:
            return {"table_type": "xlsx", "sheets": []}

        namespace = {
            "main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
        }
        sheets_payload: List[Dict[str, Any]] = []
        for sheet_name in sheet_names:
            root = ET.fromstring(archive.read(sheet_name))
            rows_payload: List[List[str]] = []
            for row in root.findall(".//main:sheetData/main:row", namespace):
                cell_values: List[str] = []
                for cell in row.findall("main:c", namespace):
                    cell_type = cell.attrib.get("t")
                    value_node = cell.find("main:v", namespace)
                    inline_node = cell.find("main:is/main:t", namespace)
                    value = ""
                    if inline_node is not None and inline_node.text:
                        value = inline_node.text
                    elif value_node is not None and value_node.text is not None:
                        raw = value_node.text
                        if cell_type == "s":
                            index = int(raw)
                            value = (
                                shared_strings[index]
                                if index < len(shared_strings)
                                else raw
                            )
                        else:
                            value = raw
                    cell_values.append(value)
                if cell_values:
                    rows_payload.append(cell_values)

            headers = rows_payload[0] if rows_payload else []
            rows = [
                {
                    headers[index]
                    if index < len(headers) and headers[index]
                    else f"column_{index + 1}": value
                    for index, value in enumerate(row)
                }
                for row in rows_payload[1:]
            ]
            sheets_payload.append(
                {
                    "sheet_name": Path(sheet_name).stem,
                    "headers": headers,
                    "rows": rows,
                }
            )

        return {"table_type": "xlsx", "sheets": sheets_payload}


def extraction_for_unsupported_visual(
    path: str, extension: str
) -> Tuple[str, Optional[str], Optional[Dict[str, Any]], Optional[str], str]:
    return (
        "pending_ocr",
        None,
        None,
        f"OCR is not enabled yet for {extension or Path(path).suffix.lower() or 'this format'}",
        "ocr_not_available",
    )


def extract_document_content(
    path: str, extension: str
) -> Tuple[str, Optional[str], Optional[Dict[str, Any]], Optional[str], str]:
    ext = extension.lower()
    if ext in TEXT_EXTENSIONS:
        return "completed", read_text_file(path), None, None, "text_reader"
    if ext in DOCX_EXTENSIONS:
        return "completed", extract_docx_text(path), None, None, "python-docx"
    if ext in PDF_EXTENSIONS:
        extracted = extract_pdf_text(path)
        if extracted.strip():
            return "completed", extracted, None, None, "pypdf"
        return extraction_for_unsupported_visual(path, ext)
    if ext in CSV_EXTENSIONS:
        table = extract_csv_table(path)
        raw_text = json.dumps(table, ensure_ascii=True)
        return "completed", raw_text, table, None, "csv"
    if ext in XLSX_EXTENSIONS:
        table = parse_xlsx_table(path)
        raw_text = json.dumps(table, ensure_ascii=True)
        return "completed", raw_text, table, None, "xlsx_zip_parser"
    if ext in XLS_EXTENSIONS:
        return (
            "unsupported",
            None,
            None,
            "XLS extraction requires a legacy parser not enabled in this phase",
            "xls_pending",
        )
    if ext in DOC_EXTENSIONS:
        return (
            "unsupported",
            None,
            None,
            "DOC extraction is not enabled yet in this phase",
            "doc_pending",
        )
    if ext in IMAGE_EXTENSIONS:
        return extraction_for_unsupported_visual(path, ext)
    return (
        "unsupported",
        None,
        None,
        f"Unsupported document extension: {ext or 'unknown'}",
        "unsupported",
    )


def build_classification_input(
    raw_text: Optional[str], extracted_table_json: Optional[Dict[str, Any]]
) -> str:
    parts: List[str] = []
    if raw_text:
        parts.append(raw_text)
    if extracted_table_json:
        parts.append(json.dumps(extracted_table_json, ensure_ascii=True))
    return "\n".join(parts)


def classify_document_content(
    raw_text: Optional[str], extracted_table_json: Optional[Dict[str, Any]]
) -> Tuple[str, float]:
    corpus = build_classification_input(raw_text, extracted_table_json)
    is_screenplay, screenplay_confidence, _signals = is_probable_screenplay(corpus)
    if is_screenplay:
        return "script", screenplay_confidence

    normalized_corpus = corpus.lower()
    best_type = "unknown_document"
    best_hits = 0
    for doc_type, keywords in DOC_TYPE_RULES.items():
        hits = sum(1 for keyword in keywords if keyword in normalized_corpus)
        if hits > best_hits:
            best_type = doc_type
            best_hits = hits
    if best_hits == 0:
        return best_type, 0.1
    confidence = min(0.35 + (best_hits * 0.2), 0.95)
    return best_type, confidence


def generate_structured_payload(
    document: DocumentAsset,
    classification: DocumentClassification,
    extraction: Optional[DocumentExtraction],
) -> Dict[str, Any]:
    raw_text = extraction.raw_text if extraction else None
    extracted_table = extraction.extracted_table_json if extraction else None
    payload: Dict[str, Any] = {
        "document_asset_id": document.id,
        "file_name": document.file_name,
        "doc_type": classification.doc_type,
        "source_kind": document.source_kind,
        "raw_text_excerpt": (raw_text or "")[:500],
        "table_preview": extracted_table,
    }

    if classification.doc_type == "sound_report":
        payload.update(
            {
                "sound_roll": None,
                "mixer_name": None,
                "boom_operator": None,
                "sample_rate": None,
                "notes": raw_text,
            }
        )
    elif classification.doc_type == "camera_report":
        payload.update(
            {
                "camera_label": None,
                "card_or_mag": None,
                "lens": None,
                "fps": None,
                "notes": raw_text,
            }
        )
    elif classification.doc_type == "script_note":
        payload.update(
            {
                "best_take": None,
                "continuity_notes": raw_text,
                "editor_note": None,
            }
        )
    elif classification.doc_type == "director_note":
        payload.update(
            {
                "preferred_take": None,
                "intention_note": raw_text,
                "pacing_note": None,
                "coverage_note": None,
            }
        )
    elif classification.doc_type == "operator_note":
        payload.update(
            {
                "operator_name": None,
                "department": None,
                "note": raw_text,
            }
        )
    elif classification.doc_type == "script":
        payload.update(
            {
                "screenplay_excerpt": (raw_text or "")[:2000],
                "manual_review_required": False,
            }
        )
    else:
        payload.update(
            {
                "summary": raw_text,
                "manual_review_required": True,
            }
        )

    return payload


async def get_owned_document(
    db: AsyncSession,
    document_id: str,
    organization_id: Optional[str],
    user_id: str,
) -> DocumentAsset:
    filters = [DocumentAsset.id == document_id, DocumentAsset.uploaded_by == user_id]
    if organization_id is not None:
        filters.append(DocumentAsset.organization_id == organization_id)
    result = await db.execute(select(DocumentAsset).where(and_(*filters)))
    document = result.scalar_one_or_none()
    if document is None:
        raise HTTPException(status_code=404, detail="Document asset not found")
    return document


async def get_latest_extraction(
    db: AsyncSession, document_id: str
) -> Optional[DocumentExtraction]:
    result = await db.execute(
        select(DocumentExtraction)
        .where(DocumentExtraction.document_asset_id == document_id)
        .order_by(desc(DocumentExtraction.created_at))
    )
    return result.scalar_one_or_none()


async def get_latest_classification(
    db: AsyncSession, document_id: str
) -> Optional[DocumentClassification]:
    result = await db.execute(
        select(DocumentClassification)
        .where(DocumentClassification.document_asset_id == document_id)
        .order_by(desc(DocumentClassification.created_at))
    )
    return result.scalar_one_or_none()


async def get_latest_structured_data(
    db: AsyncSession, document_id: str
) -> Optional[DocumentStructuredData]:
    result = await db.execute(
        select(DocumentStructuredData)
        .where(DocumentStructuredData.document_asset_id == document_id)
        .order_by(desc(DocumentStructuredData.created_at))
    )
    return result.scalar_one_or_none()


async def resolve_document_context(
    db: AsyncSession,
    payload: Dict[str, Any],
    user_id: str,
    requested_org_id: Optional[str],
    project_id: str,
    media_asset_id: Optional[str],
    storage_source_id: Optional[str],
) -> Tuple[str, Optional[MediaAsset], Optional[StorageSource]]:
    organization_id = resolve_organization_id(payload, requested_org_id, user_id)
    media_asset = None
    if media_asset_id:
        media_asset = await get_owned_asset(
            db, media_asset_id, organization_id, user_id
        )
        if media_asset.project_id != project_id:
            raise HTTPException(status_code=400, detail="Media asset project mismatch")
        organization_id = media_asset.organization_id
    source = None
    if storage_source_id:
        source = await get_owned_storage_source(
            db, storage_source_id, organization_id, user_id
        )
        if source.project_id != project_id:
            raise HTTPException(
                status_code=400, detail="Storage source project mismatch"
            )
        organization_id = source.organization_id
    return organization_id, media_asset, source


async def get_document_events(
    db: AsyncSession,
    document: DocumentAsset,
) -> List[IngestEvent]:
    result = await db.execute(
        select(IngestEvent).where(
            and_(
                IngestEvent.organization_id == document.organization_id,
                IngestEvent.project_id == document.project_id,
            )
        )
    )
    events = result.scalars().all()
    document_events: List[IngestEvent] = []
    for event in events:
        payload = event.event_payload_json or {}
        if payload.get("document_asset_id") == document.id:
            document_events.append(event)
    document_events.sort(key=lambda event: event.created_at)
    return document_events
